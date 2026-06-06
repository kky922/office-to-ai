"""Word (.docx) reader using python-docx. Preserves headings, lists, tables, images."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from .common import Block, ImageAsset


def _style_heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    s = style_name.strip().lower()
    if s.startswith("heading "):
        try:
            return max(1, min(6, int(s.split()[-1])))
        except ValueError:
            return None
    if s in {"title"}:
        return 1
    if s in {"subtitle"}:
        return 2
    return None


def _list_prefix(paragraph) -> str | None:
    # A paragraph is a list item when either (a) its style is a List* style,
    # or (b) its XML has a w:numPr element (direct numbering).
    style_name = (paragraph.style.name if paragraph.style else "") or ""
    s = style_name.lower()
    if s.startswith("list "):
        return "- "
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl_el = numPr.find(qn("w:ilvl"))
    lvl = 0
    if ilvl_el is not None:
        try:
            lvl = int(ilvl_el.get(qn("w:val")) or 0)
        except ValueError:
            lvl = 0
    return "  " * lvl + "- "


def _inline(text: str) -> str:
    return text.replace("|", "\\|")


def _paragraph_md(paragraph) -> str:
    text = paragraph.text or ""
    if not text.strip():
        return ""
    level = _style_heading_level(paragraph.style.name if paragraph.style else None)
    if level is not None:
        return "#" * level + " " + text.strip()
    prefix = _list_prefix(paragraph)
    if prefix is not None:
        return prefix + text.strip()
    return text.strip()


def _table_md(table) -> str:
    rows = [[_inline(cell.text.strip()) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    head = "| " + " | ".join(rows[0]) + " |"
    sep = "|" + "|".join(" --- " for _ in range(width)) + "|"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    return "\n".join(filter(None, [head, sep, body]))


def read(path: str | Path, image_mode: str = "extract") -> tuple[list[Block], list[ImageAsset]]:
    path = Path(path)
    doc = Document(str(path))
    blocks: list[Block] = [
        Block(
            kind="heading",
            md=f"# {path.name}\n\n> Word 문서\n",
            source=path.name,
        )
    ]
    images: list[ImageAsset] = []
    current_heading = path.name
    img_counter = 0

    # Iterate in document order by walking the body's child elements.
    body = doc.element.body
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    drawing_tag = qn("w:drawing")
    blip_tag = qn("a:blip")

    p_map = {p._p: p for p in doc.paragraphs}
    t_map = {t._tbl: t for t in doc.tables}

    for child in body.iterchildren():
        if child.tag == p_tag:
            para = p_map.get(child)
            if para is None:
                continue
            md = _paragraph_md(para)
            if md:
                if md.startswith("#"):
                    current_heading = md.lstrip("# ").strip() or current_heading
                    blocks.append(Block(kind="heading", md=md, source=current_heading))
                else:
                    blocks.append(Block(kind="paragraph", md=md, source=current_heading))
            # inline drawings → image references
            for drawing in child.iter(drawing_tag):
                for blip in drawing.iter(blip_tag):
                    rid = blip.get(qn("r:embed"))
                    if not rid:
                        continue
                    try:
                        part = doc.part.related_parts[rid]
                    except KeyError:
                        continue
                    img_counter += 1
                    if image_mode == "omit":
                        continue
                    if image_mode == "placeholder":
                        blocks.append(
                            Block(
                                kind="paragraph",
                                md=f"_[이미지 {img_counter}]_",
                                source=current_heading,
                            )
                        )
                        continue
                    ext = (part.partname.split(".")[-1] or "bin").lower()
                    name = f"word_img_{img_counter}.{ext}"
                    images.append(ImageAsset(path=f"images/{name}", data=part.blob))
                    blocks.append(
                        Block(
                            kind="paragraph",
                            md=f"![이미지 {img_counter}](images/{name})",
                            source=current_heading,
                        )
                    )
        elif child.tag == tbl_tag:
            tbl = t_map.get(child)
            if tbl is None:
                continue
            md = _table_md(tbl)
            if md:
                blocks.append(Block(kind="table", md=md + "\n", source=current_heading))
    return blocks, images
